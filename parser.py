import re
import os
from datetime import datetime, timedelta
import docx
from docx import Document
import threading

qa_progress = []  # live Q&A for web

# Timer class left unchanged
class ExamTimer(threading.Thread):
    def __init__(self, duration_seconds=7200):
        super().__init__(daemon=True)
        self.start_time = datetime.now()
        self.end_time = self.start_time + timedelta(seconds=duration_seconds)

    def remaining_time(self):
        now = datetime.now()
        if now >= self.end_time:
            return 0
        return int((self.end_time - now).total_seconds())

    def formatted_remaining(self):
        seconds = self.remaining_time()
        if seconds <= 0:
            return "Time is up."
        mins, secs = divmod(seconds, 60)
        hrs, mins = divmod(mins, 60)
        return f"{hrs} hours {mins} minutes {secs} seconds remaining"


# Your config vars and excluded substrings unchanged
INPUT_DOC = "mugilanQp.docx"
OUTPUT_DOC = "answers.docx"


EXCLUDE_SUBSTRS = {
    "answer all questions",
    "name & signature",
    "department of data science",
    "max. marks",
    "time duration",
    "affiliated to",
    "college",
    "batch:", "class:", "subject title:", "semester:",
    "mid term", "reviewer"
}

# No whisper or local recording imports


def get_all_text(doc_path):
    doc = docx.Document(doc_path)
    lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    lines.append(txt)
    return lines


def clean_line(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip())


def is_excluded(line: str) -> bool:
    l = line.lower()
    return any(key in l for key in EXCLUDE_SUBSTRS)


def extract_questions(path: str):
    raw_lines = get_all_text(path)
    lines, prev = [], None
    for l in raw_lines:
        l = clean_line(l)
        if l and l != prev:
            lines.append(l)
        prev = l

    section, result, i = None, [], 0
    while i < len(lines):
        line = lines[i]
        if line.lower().startswith("section a"):
            section = "A"
            i += 1
            continue
        if line.lower().startswith("section b"):
            section = "B"
            i += 1
            continue
        if line.lower().startswith("section c"):
            section = "C"
            i += 1
            continue

        if not section or is_excluded(line):
            i += 1
            continue

        # Section A - MCQ questions (unchanged)
        if section == "A":
            m = re.match(r"^(\d{1,2})\s+(.*)$", line)
            if m:
                qnum, stem_text = m.groups()
                i += 1
            elif re.fullmatch(r"\d{1,2}", line):
                qnum = line
                i += 1
                stem_text = ""
                if i < len(lines) and not re.match(r"^[ABCD]\b", lines[i], re.I):
                    stem_text = lines[i].strip()
                    i += 1
            else:
                i += 1
                continue

            options = {}
            while i < len(lines) and re.match(r"^[ABCD]\b", lines[i], re.I):
                letter = lines[i][0].upper()
                value = lines[i][1:].strip()
                options[letter] = value
                i += 1

            text = stem_text
            for letter in ["A", "B", "C", "D"]:
                if letter in options:
                    text += f"\n{letter}. {options[letter]}"

            result.append({"section": "A", "label": qnum, "text": text})
            continue

        # Section B/C questions (unchanged)
        m = re.match(r"^(\d{1,2})\s*([AB])?\s*(.*)$", line)
        if section in {"B", "C"} and m:
            qnum, ab, rest = m.groups()
            ab = ab or ""
            block = [rest] if rest else []
            i += 1
            while i < len(lines) and not re.match(r"^\d{1,2}\s*[AB]?", lines[i]) and not lines[i].lower().startswith("section"):
                if not is_excluded(lines[i]) and lines[i] != "(OR)":
                    if lines[i] in {"A", "B"} and not block:
                        ab = lines[i]
                    else:
                        block.append(lines[i])
                i += 1
            text = " ".join(block)
            result.append({"section": section, "label": f"{qnum} {ab}".strip(), "text": text})
            continue

        i += 1

    return result


def save_answers_docx(path: str, qa_items):
    doc = Document()
    doc.add_heading("Answers Document", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    for item in qa_items:
        doc.add_paragraph(f"Q{item['label']}: {item['text']}", style="List Bullet")
        doc.add_paragraph(f"A{item['label']}: {item.get('answer', '').strip()}\n")
    doc.save(path)


def extract_metadata(path: str):
    raw_lines = get_all_text(path)
    name, subject = None, None
    for line in raw_lines:
        if line.strip().lower().startswith("name:"):
            name = line.split(":", 1)[-1].strip()
        if line.strip().lower().startswith("subject title:"):
            subject = line.split(":", 1)[-1].strip()
    return name, subject


def main():
    # This main will not do local audio recording or whisper transcription
    # The transcription and answers must come from the web frontend or other input

    timer = ExamTimer(duration_seconds=7200)
    timer.start()

    if not os.path.exists(INPUT_DOC):
        raise FileNotFoundError(f"Input DOCX not found: {INPUT_DOC}")

    name, subject = extract_metadata(INPUT_DOC)
    if not name or not subject:
        print("⚠️ Could not find name/subject in the paper.")
        return

    print(f"📄 Candidate: {name}, Subject: {subject}")

    # You can implement logic here to wait for answers to come via frontend API or other method.
    # For example, this could poll a database, queue, or shared variable updated by frontend endpoint.

    # For demonstration, just load questions without answers
    questions = extract_questions(INPUT_DOC)
    print(f"✅ Extracted {len(questions)} questions.")

    # Here instead of recording and transcribing, you await answers from frontend
    # For testing, mark all answers skipped
    qa_items = [{"label": q["label"], "text": q["text"], "answer": "[Pending answer via Web]"} for q in questions]

    # Save placeholder answers for now
    save_answers_docx(OUTPUT_DOC, qa_items)
    print(f"✅ Saved placeholder answers to {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
